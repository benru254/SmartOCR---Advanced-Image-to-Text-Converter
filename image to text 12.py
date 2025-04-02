import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
from langdetect import detect
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class SmartOCRApp:
    def __init__(self, root):
        """
        Initialize the Smart OCR Application with improved text filtering and multi-image support.
        """
        # Ensure Tesseract is installed
        tesseract_path = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
        if not os.path.exists(tesseract_path):
            messagebox.showerror("Tesseract Not Found", "Tesseract is not installed or misconfigured.")
            root.quit()
        pytesseract.pytesseract.tesseract_cmd = tesseract_path

        # Configure root window
        self.root = root
        self.root.title("Smart OCR - Multi-Image Extraction")
        self.root.geometry("900x700")
        self.root.configure(bg='#f4f4f4')

        # Store extracted texts with image sources
        self.extracted_texts = {}

        # Create UI components
        self.create_interface()

    def create_interface(self):
        """Create the enhanced user interface."""
        frame = tk.Frame(self.root, bg='#f4f4f4')
        frame.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)

        # Top Button Frame
        top_button_frame = tk.Frame(frame, bg='#f4f4f4')
        top_button_frame.pack(fill=tk.X, pady=(0, 10))

        # Select Multiple Images Button
        btn_open = tk.Button(
            top_button_frame, text="Select Images", command=self.open_images,
            bg='#2C3E50', fg='white', padx=15, pady=8, font=('Arial', 10, 'bold')
        )
        btn_open.pack(side=tk.LEFT, padx=(0, 10))

        # Export Dropdown
        export_label = tk.Label(top_button_frame, text="Export As:", bg='#f4f4f4')
        export_label.pack(side=tk.LEFT, padx=(10, 5))

        self.export_var = tk.StringVar(value="Text")
        export_options = ["Text", "Word", "PDF"]
        export_dropdown = ttk.Combobox(
            top_button_frame, 
            textvariable=self.export_var, 
            values=export_options, 
            state="readonly", 
            width=10
        )
        export_dropdown.pack(side=tk.LEFT, padx=(0, 10))

        # Save Text Button
        btn_save = tk.Button(
            top_button_frame, text="Export Extracted Text", command=self.save_extracted_text,
            bg='#34495E', fg='white', padx=15, pady=8, font=('Arial', 10, 'bold')
        )
        btn_save.pack(side=tk.LEFT)

        # Image List Frame
        self.image_list_frame = tk.Frame(frame, bg='#f4f4f4')
        self.image_list_frame.pack(fill=tk.X, pady=(0, 10))

        # Scrollable Text Display
        self.text_display = scrolledtext.ScrolledText(
            frame, wrap=tk.WORD, height=20, width=80,
            font=('Consolas', 10), bg='white', borderwidth=2, relief=tk.SUNKEN
        )
        self.text_display.pack(fill=tk.BOTH, expand=True, pady=10)

    def open_images(self):
        """Open file dialog to select multiple images."""
        file_paths = filedialog.askopenfilenames(filetypes=[
            ("Image Files", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff"),
            ("All Files", "*.*")
        ])
        
        if file_paths:
            # Clear previous extraction
            self.extracted_texts.clear()
            self.text_display.delete(1.0, tk.END)
            
            # Clear previous image list
            for widget in self.image_list_frame.winfo_children():
                widget.destroy()
            
            # Process each selected image
            for file_path in file_paths:
                # Create label for each image
                filename = os.path.basename(file_path)
                img_label = tk.Label(
                    self.image_list_frame, 
                    text=filename, 
                    bg='#f4f4f4', 
                    anchor='w'
                )
                img_label.pack(fill=tk.X)
                
                # Extract text
                extracted_text = self.extract_text_from_image(file_path)
                
                # Store extracted text with image source
                self.extracted_texts[filename] = extracted_text
                
                # Append to text display
                self.text_display.insert(tk.END, f"--- {filename} ---\n")
                self.text_display.insert(tk.END, extracted_text + "\n\n")

            messagebox.showinfo("Success", f"Extracted text from {len(file_paths)} images.")

    def preprocess_image(self, image_path):
        """
        Preprocess the image for better OCR accuracy:
        - Convert to grayscale
        - Increase contrast
        - Sharpen text
        - Remove noise
        """
        image = Image.open(image_path)
        image = image.convert("L")  # Convert to grayscale
        image = image.filter(ImageFilter.MedianFilter())  # Reduce noise
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2)  # Increase contrast
        return image

    def filter_unreadable_text(self, text):
        """
        Remove symbols, gibberish, and unreadable characters while keeping readable text.
        """
        # Keep only words with valid English letters and common punctuation
        filtered_text = re.sub(r'[^A-Za-z0-9.,!?;:\'"\-\s]', '', text)
        return filtered_text

    def extract_text_from_image(self, image_path):
        """
        Extract and clean text from the selected image.
        """
        try:
            # Preprocess image
            image = self.preprocess_image(image_path)

            # Extract text using OCR
            raw_text = pytesseract.image_to_string(image, lang="eng")

            # Remove unreadable characters
            cleaned_text = self.filter_unreadable_text(raw_text)

            return cleaned_text

        except Exception as e:
            messagebox.showerror("OCR Error", f"Could not process image:\n{str(e)}")
            return ""

    def save_extracted_text(self):
        """Save extracted text in selected format."""
        if not self.extracted_texts:
            messagebox.showwarning("Empty Text", "No text available to save.")
            return

        # Determine save path and format
        export_format = self.export_var.get().lower()
        file_types = {
            "text": [("Text Files", "*.txt")],
            "word": [("Word Documents", "*.docx")],
            "pdf": [("PDF Files", "*.pdf")]
        }
        
        save_path = filedialog.asksaveasfilename(
            defaultextension=f".{export_format}",
            filetypes=file_types[export_format]
        )

        if save_path:
            try:
                # Generate output based on selected format
                if export_format == "text":
                    self._save_as_text(save_path)
                elif export_format == "word":
                    self._save_as_word(save_path)
                elif export_format == "pdf":
                    self._save_as_pdf(save_path)
                
                messagebox.showinfo("Saved", f"Text saved successfully to:\n{save_path}")
            except Exception as e:
                messagebox.showerror("Save Error", f"Could not save file:\n{str(e)}")

    def _save_as_text(self, save_path):
        """Save extracted texts to a text file."""
        with open(save_path, 'w', encoding='utf-8') as file:
            for image_name, text in self.extracted_texts.items():
                file.write(f"--- {image_name} ---\n")
                file.write(text + "\n\n")

    def _save_as_word(self, save_path):
        """Save extracted texts to a Word document."""
        doc = docx.Document()
        for image_name, text in self.extracted_texts.items():
            doc.add_heading(image_name, level=2)
            doc.add_paragraph(text)
            doc.add_paragraph()  # Add blank line between entries
        doc.save(save_path)

    def _save_as_pdf(self, save_path):
        """Save extracted texts to a PDF."""
        c = canvas.Canvas(save_path, pagesize=letter)
        width, height = letter
        y = height - 100  # Start near the top of the page
        
        c.setFont("Helvetica", 12)
        for image_name, text in self.extracted_texts.items():
            # Add image name as a header
            c.drawString(100, y, f"--- {image_name} ---")
            y -= 20
            
            # Split text into lines
            lines = text.split('\n')
            for line in lines:
                # Move to next page if near bottom
                if y < 100:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = height - 100
                
                # Draw text line
                c.drawString(100, y, line)
                y -= 15
            
            # Add some space between entries
            y -= 30
        
        c.save()


def main():
    """Run the application."""
    root = tk.Tk()
    app = SmartOCRApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()